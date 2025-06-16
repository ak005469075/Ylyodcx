import os
import re
from docx import Document
from docx.oxml.shared import qn
from docx.oxml import OxmlElement
from tips import utils
import io

#By yLy
def handle_docx(data,Myconfig):

    heading_pos=Myconfig["head_pos"]
    head_pos_end=Myconfig["head_pos_end"]
    init_heading_new_pos=Myconfig["head_new_pos"]
    
    #print(docx_path)

   
    doc = Document(data)
    
    utils.remove_konghang(doc)
    new_doc=Document("muban.docx")

##    for i,para in enumerate(doc.paragraphs):
##        print(i)
##        print(para.text)
        

    work_order_number = "123"
    for para in doc.paragraphs:
        match = re.search(r'[A-Z]+\d+', para.text)
        if match:
            work_order_number = match.group(0)
            print(f"单号{work_order_number}")
            break
    


    #先正文后标题，防止index变动影响
                  

    #标题预览
    print("标题预览：")
    old_heads=OxmlElement('w:body')
    for para in doc.paragraphs[heading_pos:heading_pos+head_pos_end]:
          print(para.text)
          
          old_heads.append(para._element.__copy__())
          
          
    #正文预览
    print("正文预览：") 
    heading_new_pos=init_heading_new_pos

    old_paras=OxmlElement('w:body')
    for para in doc.paragraphs[heading_pos+head_pos_end:]:
          print(para.text)
          old_paras.append(para._element.__copy__())
          
    
          
    #先复制正文
          
    target_paragraph=new_doc.paragraphs[heading_new_pos+1]
    #print(target_paragraph.text)
    parent = target_paragraph._element.getparent()
    #print(target_paragraph._element)
    parent.remove(target_paragraph._element)
    for element in old_paras:
        heading_new_pos+=1
        parent.insert(heading_new_pos,element.__copy__())


    
          
    #后复制标题
          
    t_head=new_doc.paragraphs[init_heading_new_pos]
    #print(t_head.text)
    parent_h = t_head._element.getparent()
    #print(t_head._element)
    parent_h.remove(t_head._element)
    for element in old_heads:
        parent_h.insert(init_heading_new_pos,element.__copy__())
        init_heading_new_pos+=1


     #写入序号
    xuhao=input("请输入序号： ")
    for para_new in new_doc.paragraphs:
          for run in para_new.runs:
              #print(run.text)
              if "xh" in run.text:
                  #print(run.text)
                  run.text = run.text.replace("xh",xuhao)
    ##           print(j)
    ##           print(run.text)
        
    return doc,new_doc,work_order_number
